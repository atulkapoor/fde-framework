"""Reading what an FDE is actually handed.

Briefs arrive as PDFs and Word documents far more often than as plain text. The
rule here is narrow and deliberate: **read what can be read properly, and refuse
the rest by name.**

Reading a PDF's bytes as text is the failure worth guarding against. It does not
raise -- it produces a string full of tokens that look like words, which the
prose parser then reads facts out of. Facts extracted from binary noise carry
artifact provenance and outrank the interview answer that would have corrected
them, which makes silent success far worse here than a clean refusal.

PDF and Word support is a missing optional dependency rather than a policy. When
a reader is installed it is used; when it is not, the message says so and points
at the thing that works right now, because an FDE holding a PDF and a deadline
needs the next step rather than a diagnosis.
"""

from __future__ import annotations

from pathlib import Path

# Formats that carry text inside a container. Reading their bytes produces
# plausible-looking rubbish, so each is refused by name until a reader exists.
NEEDS_A_READER = {
    ".pdf": ("pypdf", "pdf"),
    ".docx": ("python-docx", "docx"),
    ".doc": ("antiword or a conversion to .docx", "doc"),
    ".pptx": ("python-pptx", "pptx"),
    ".xlsx": ("openpyxl", "xlsx"),
}

READS_AS_TEXT = {".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".yaml", ".yml", ""}


class UnreadableDocument(Exception):
    """This file cannot be turned into text safely."""


def read_document(path: str | Path) -> str:
    """Text from a file, or a refusal that says what to do instead."""
    path = Path(path)

    if not path.exists():
        raise UnreadableDocument(f"{path} does not exist")

    suffix = path.suffix.lower()

    if suffix in NEEDS_A_READER:
        package, name = NEEDS_A_READER[suffix]
        try:
            extracted = _extract(suffix, path)
        except Exception as exc:  # noqa: BLE001 - each reader raises its own kinds
            # The module's whole contract is refusal by name. A corrupt file
            # escaping as a pypdf traceback breaks it at the one moment it
            # was needed.
            raise UnreadableDocument(
                f"{path.name}: the {name} reader could not read it -- {exc}"
            ) from exc
        if extracted is not None:
            return _require_content(extracted, path)
        raise UnreadableDocument(
            f"{path.name} is a {name} file and no reader for it is installed. "
            f"Install {package}, or paste the text with --text. "
            f"Reading it as bytes would produce facts from noise, which is worse "
            f"than reading nothing."
        )

    if suffix not in READS_AS_TEXT:
        raise UnreadableDocument(
            f"{path.name}: unrecognised extension {suffix!r}. If it is plain text, "
            f"rename it or paste it with --text."
        )

    try:
        # Never strict. A brief with a stray byte is still a brief, and failing
        # on one character would be its own kind of unhelpful.
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise UnreadableDocument(f"{path}: {exc}") from exc

    return _require_content(text, path)


def _require_content(text: str, path: Path) -> str:
    """An empty file is reported rather than passed on.

    Nothing extracted from an empty file looks exactly like nothing extracted
    from a full one the parser could not understand, and those need different
    responses.
    """
    if not text.strip():
        raise UnreadableDocument(f"{path.name} is empty")
    return text


def _extract(suffix: str, path: Path) -> str | None:
    """Delegate to a reader if one is available. None means there is not one."""
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    return None


def _extract_pdf(path: Path) -> str | None:
    try:
        from pypdf import PdfReader
    except ImportError:
        return None
    return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)


def _extract_docx(path: Path) -> str | None:
    try:
        import docx
    except ImportError:
        return None
    return "\n".join(p.text for p in docx.Document(str(path)).paragraphs)
